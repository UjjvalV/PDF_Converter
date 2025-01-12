import streamlit as st
from docx import Document

def extract_metadata(doc):
    metadata = {
        "Number of paragraphs": len(doc.paragraphs),
        "Number of tables": len(doc.tables),
        "Word count": sum(len(paragraph.text.split()) for paragraph in doc.paragraphs),
    }
    return metadata

def upload_and_display_metadata():
    uploaded_file = st.file_uploader("Upload a DOCX file", type=["docx"])
    
    if uploaded_file:
        doc = Document(uploaded_file)
        metadata = extract_metadata(doc)
        
        st.subheader("Metadata:")
        for key, value in metadata.items():
            st.write(f"**{key}:** {value}")
        
        return uploaded_file
    
    return None
